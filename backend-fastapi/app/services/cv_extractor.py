from __future__ import annotations

import logging
import os

import fitz

logger = logging.getLogger(__name__)


def extract_text_from_pdf(pdf_bytes: bytes) -> str:
    """Extract text from a PDF file using PyMuPDF (fitz)."""
    if not pdf_bytes:
        raise ValueError("Could not download CV file")

    logger.info("PDF extraction: received %d bytes", len(pdf_bytes))

    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as exc:  # pragma: no cover - library-specific parse errors
        logger.error("fitz.open() failed: %s", exc)
        raise ValueError("Could not parse CV file") from exc

    logger.info("PDF has %d page(s)", len(doc))

    chunks: list[str] = []
    for i, page in enumerate(doc):
        page_text = page.get_text("text")
        logger.info("Page %d: extracted %d chars", i + 1, len(page_text))
        chunks.append(page_text)

    full_text = "\n".join(chunks)
    cleaned = full_text.replace(" \n", "\n").strip()

    while "\n\n\n" in cleaned:
        cleaned = cleaned.replace("\n\n\n", "\n\n")

    if not cleaned:
        # Log the first 200 raw bytes to help diagnose what was downloaded
        logger.error(
            "PDF text empty after extraction. First 200 raw bytes: %r",
            pdf_bytes[:200],
        )
        raise ValueError("CV text is empty after extraction")

    logger.info("PDF extraction successful: %d chars", len(cleaned))
    return cleaned


def extract_text_from_docx(docx_bytes: bytes) -> str:
    """Extract text from a Word (.docx) file using python-docx."""
    if not docx_bytes:
        raise ValueError("Could not download CV file")

    logger.info("DOCX extraction: received %d bytes", len(docx_bytes))

    try:
        from io import BytesIO

        from docx import Document  # lazy import — python-docx may not be installed

        doc = Document(BytesIO(docx_bytes))
    except ImportError:
        raise ValueError(
            "Word document support is not available on this server. Please upload a PDF instead."
        )
    except Exception as exc:  # pragma: no cover - library-specific parse errors
        logger.error("python-docx parsing failed: %s", exc)
        raise ValueError("Could not parse CV file") from exc

    chunks: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    # Also extract text from tables (common in CV templates)
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                chunks.append("  ".join(row_texts))

    full_text = "\n".join(chunks)
    cleaned = full_text.replace(" \n", "\n").strip()

    while "\n\n\n" in cleaned:
        cleaned = cleaned.replace("\n\n\n", "\n\n")

    if not cleaned:
        logger.error(
            "DOCX text empty after extraction. First 200 raw bytes: %r",
            docx_bytes[:200],
        )
        raise ValueError("CV text is empty after extraction")

    logger.info("DOCX extraction successful: %d chars", len(cleaned))
    return cleaned


def extract_cv_text(file_bytes: bytes, filename: str) -> str:
    """Route to the correct extractor based on file extension.

    Raises ValueError for unsupported file types or extraction failures.
    """
    ext = os.path.splitext(filename.lower())[1]
    logger.info("extract_cv_text called: filename=%r, ext=%r, %d bytes", filename, ext, len(file_bytes))

    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext == ".docx":
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(
            f"Unsupported file type '{ext}'. Please upload a PDF or Word (.docx) file."
        )
